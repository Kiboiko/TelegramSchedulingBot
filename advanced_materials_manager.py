# advanced_materials_manager.py
import logging
import os
from datetime import datetime
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Inches
import asyncio

# Импортируем класс из DocsMerge.py
from DocsMerge import GoogleDocsMerger

logger = logging.getLogger(__name__)


class AdvancedMaterialsManager:
    def __init__(self, gsheets_manager, credentials_file: str, spreadsheet_id: str):
        self.gsheets = gsheets_manager
        self.credentials_file = credentials_file
        self.spreadsheet_id = spreadsheet_id
        self.output_dir = "combined_materials"
        self.docs_merger = GoogleDocsMerger(credentials_file)

        # Создаем папку для документов если её нет
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
            logger.info(f"Создана папка для документов: {self.output_dir}")

    def get_student_lesson_info(self, user_id: int, subject_id: str, target_date: str) -> Dict[str, any]:
        """Получает информацию о занятии ученика включая номер занятия"""
        try:
            from shedule_app.GoogleParser import GoogleSheetsDataLoader

            loader = GoogleSheetsDataLoader(self.credentials_file, self.spreadsheet_id, target_date)

            # Получаем номер занятия для ученика
            lesson_number = loader.get_lesson_number_for_student(str(user_id), target_date, str(subject_id))

            # Получаем тему занятия
            topic = loader.get_student_topic_by_user_id(str(user_id), target_date, str(subject_id))

            return {
                'lesson_number': lesson_number or 1,
                'topic': topic or "Тема не определена",
                'subject_name': self._get_subject_name(subject_id)
            }
        except Exception as e:
            logger.error(f"Ошибка получения информации о занятии: {e}")
            return {
                'lesson_number': 1,
                'topic': "Тема не определена",
                'subject_name': self._get_subject_name(subject_id)
            }

    def get_document_links_for_subjects(self) -> Dict[str, str]:
        """Получает все ссылки на документы из листа квалификаций"""
        try:
            document_links = {}

            # Используем существующий функционал из gsheets_manager
            if hasattr(self.gsheets, 'qual_links'):
                document_links = self.gsheets.qual_links.copy()
                logger.info(f"Найдено {len(document_links)} ссылок на документы")
            else:
                logger.warning("qual_links не найден в gsheets_manager")

            return document_links

        except Exception as e:
            logger.error(f"Ошибка получения ссылок на документы: {e}")
            return {}

    def get_students_with_lessons(self, target_date: str) -> List[Dict]:
        """Получает список студентов с их занятиями на указанную дату"""
        try:
            students_data = []

            # Получаем данные из листа учеников
            worksheet = self.gsheets._get_or_create_worksheet("Ученики бот")
            data = worksheet.get_all_values()

            if len(data) < 2:
                logger.warning("В листе 'Ученики бот' нет данных")
                return students_data

            # Находим колонку с выбранной датой
            date_columns = self._find_date_columns(data, target_date)
            if date_columns == (-1, -1):
                logger.error(f"Дата {target_date} не найдена в расписании")
                return students_data

            start_col, end_col = date_columns

            # Обрабатываем каждого студента
            for row in data[1:]:  # Пропускаем заголовок
                if not row or len(row) <= max(start_col, end_col):
                    continue

                user_id = str(row[0]).strip() if row[0] else ""
                student_name = str(row[1]).strip() if len(row) > 1 else ""
                subject_id = str(row[2]).strip() if len(row) > 2 else ""

                # Проверяем, есть ли запись на указанную дату
                start_time = str(row[start_col]).strip() if len(row) > start_col and row[start_col] else ""
                end_time = str(row[end_col]).strip() if len(row) > end_col and row[end_col] else ""

                if user_id and student_name and subject_id and start_time and end_time:
                    # Получаем информацию о занятии
                    lesson_info = self.get_student_lesson_info(int(user_id), subject_id, target_date)

                    students_data.append({
                        'user_id': user_id,
                        'name': student_name,
                        'subject_id': subject_id,
                        'subject_name': lesson_info['subject_name'],
                        'lesson_number': lesson_info['lesson_number'],
                        'topic': lesson_info['topic'],
                        'start_time': start_time,
                        'end_time': end_time
                    })

            logger.info(f"Найдено {len(students_data)} студентов с занятиями на {target_date}")
            return students_data

        except Exception as e:
            logger.error(f"Ошибка получения студентов с занятиями: {e}")
            return []

    def _find_date_columns(self, data: List[List], target_date: str) -> tuple:
        """Находит колонки с указанной датой в заголовках"""
        try:
            if not data:
                return (-1, -1)

            headers = data[0]
            formatted_date = self._format_date_for_search(target_date)

            for i, header in enumerate(headers):
                header_str = str(header).lower().strip()
                if formatted_date in header_str:
                    # Предполагаем, что следующая колонка - время окончания
                    if i + 1 < len(headers):
                        return (i, i + 1)
                    else:
                        return (i, i)

            return (-1, -1)

        except Exception as e:
            logger.error(f"Ошибка поиска колонок даты: {e}")
            return (-1, -1)

    def _format_date_for_search(self, date_str: str) -> str:
        """Форматирует дату для поиска в заголовках"""
        try:
            # Пробуем разные форматы дат
            date_formats = ['%d.%m.%Y', '%Y-%m-%d', '%d.%m']
            date_obj = None

            for fmt in date_formats:
                try:
                    date_obj = datetime.strptime(date_str, fmt)
                    break
                except ValueError:
                    continue

            if date_obj:
                return date_obj.strftime('%d.%m.%Y').lower()
            else:
                return date_str.lower()

        except Exception as e:
            logger.error(f"Ошибка форматирования даты для поиска: {e}")
            return date_str.lower()

    def create_combined_qualification_document(self, target_date: str) -> str:
        """Создает объединенный документ со всеми материалами по квалификациям"""
        try:
            logger.info(f"Начинаем создание объединенного документа для даты {target_date}")

            # 1. Получаем ссылки на документы по предметам
            document_links = self.get_document_links_for_subjects()
            if not document_links:
                return "❌ Не найдено ссылок на документы в квалификациях"

            logger.info(f"Найдено {len(document_links)} документов для объединения")

            # 2. Получаем студентов с занятиями на указанную дату
            students_data = self.get_students_with_lessons(target_date)
            if not students_data:
                return f"❌ Не найдено занятий на дату {target_date}"

            logger.info(f"Найдено {len(students_data)} занятий на дату {target_date}")

            # 3. Собираем уникальные предметы, которые есть в расписании
            scheduled_subjects = set()
            subject_students_map = {}

            for student in students_data:
                subject_id = student['subject_id']
                scheduled_subjects.add(subject_id)

                if subject_id not in subject_students_map:
                    subject_students_map[subject_id] = []
                subject_students_map[subject_id].append(student)

            # 4. Фильтруем документы только для тех предметов, которые есть в расписании
            relevant_document_urls = []
            subject_names_map = {}

            for subject_id, doc_url in document_links.items():
                if subject_id in scheduled_subjects:
                    relevant_document_urls.append(doc_url)
                    subject_names_map[subject_id] = self._get_subject_name(subject_id)
                    logger.info(f"Добавлен документ для предмета {subject_id}: {doc_url}")

            if not relevant_document_urls:
                return "❌ Нет документов для предметов, запланированных на эту дату"

            logger.info(f"Будут объединены {len(relevant_document_urls)} документов")

            # 5. Используем GoogleDocsMerger для объединения документов
            output_filename = f"combined_materials_{target_date.replace('.', '_')}.docx"
            output_path = os.path.join(self.output_dir, output_filename)

            # Создаем временный документ с информацией о студентах
            temp_doc_path = self._create_students_info_document(students_data, target_date)

            # Объединяем все документы
            success = self.docs_merger.merge_documents_with_images(
                relevant_document_urls,
                output_path
            )

            if success:
                # Добавляем информацию о студентах в начало объединенного документа
                final_doc = self._combine_documents(temp_doc_path, output_path, target_date)

                # Очищаем временные файлы
                if os.path.exists(temp_doc_path):
                    os.remove(temp_doc_path)

                abs_path = os.path.abspath(final_doc)
                logger.info(f"Объединенный документ создан: {abs_path}")

                student_count = len(students_data)
                doc_count = len(relevant_document_urls)

                return (f"✅ Объединенный документ создан!\n"
                        f"📁 {abs_path}\n\n"
                        f"📊 Статистика:\n"
                        f"• Студентов: {student_count}\n"
                        f"• Документов: {doc_count}\n"
                        f"• Дата: {target_date}")
            else:
                return "❌ Не удалось объединить документы"

        except Exception as e:
            logger.error(f"Ошибка создания объединенного документа: {e}")
            return f"❌ Ошибка: {str(e)}"

    def _create_students_info_document(self, students_data: List[Dict], target_date: str) -> str:
        """Создает временный документ с информацией о студентах БЕЗ ЗАГОЛОВКОВ И РАЗРЫВОВ"""
        try:
            doc = Document()

            # Группируем по предметам
            subjects_map = {}
            for student in students_data:
                subject_id = student['subject_id']
                if subject_id not in subjects_map:
                    subjects_map[subject_id] = []
                subjects_map[subject_id].append(student)

            # Добавляем информацию по каждому предмету без заголовков
            for subject_id, students in subjects_map.items():
                for student in students:
                    # Добавляем информацию о студенте как обычный параграф
                    p = doc.add_paragraph()
                    p.add_run(
                        f"{student['name']} - Занятие №{student['lesson_number']}: "
                        f"{student['topic']} ({student['start_time']}-{student['end_time']})"
                    )

            # Сохраняем временный файл
            temp_path = os.path.join(self.output_dir, f"temp_students_info_{datetime.now().timestamp()}.docx")
            doc.save(temp_path)

            return temp_path

        except Exception as e:
            logger.error(f"Ошибка создания документа с информацией о студентах: {e}")
            temp_path = os.path.join(self.output_dir, f"temp_empty_{datetime.now().timestamp()}.docx")
            doc = Document()
            doc.save(temp_path)
            return temp_path

    def _combine_documents(self, first_doc_path: str, second_doc_path: str, target_date: str) -> str:
        """Объединяет два документа в один БЕЗ РАЗДЕЛИТЕЛЕЙ И РАЗРЫВОВ СТРАНИЦ"""
        try:
            final_filename = f"final_combined_materials_{target_date.replace('.', '_')}.docx"
            final_path = os.path.join(self.output_dir, final_filename)

            # Создаем новый документ
            final_doc = Document()

            # Копируем содержимое первого документа (информация о студентах)
            first_doc = Document(first_doc_path)
            for element in first_doc.element.body:
                final_doc.element.body.append(element)

            # Копируем содержимое второго документа (материалы по предметам)
            # Без каких-либо разделителей - содержимое продолжается
            second_doc = Document(second_doc_path)
            for element in second_doc.element.body:
                final_doc.element.body.append(element)

            # Сохраняем финальный документ
            final_doc.save(final_path)

            return final_path

        except Exception as e:
            logger.error(f"Ошибка объединения документов: {e}")
            return second_doc_path

    def _get_subject_name(self, subject_id: str) -> str:
        """Получает название предмета по ID"""
        from config import SUBJECTS
        return SUBJECTS.get(subject_id, f"Предмет {subject_id}")

    async def create_combined_materials_document(self, target_date: str) -> str:
        """Асинхронная версия метода для использования в боте"""
        try:
            # Запускаем синхронный метод в отдельном потоке
            result = await asyncio.to_thread(
                self.create_combined_qualification_document,
                target_date
            )
            return result
        except Exception as e:
            logger.error(f"Ошибка в асинхронном методе: {e}")
            return f"❌ Ошибка: {str(e)}"