# documents_merger.py
import logging
from datetime import datetime
import os
from typing import Dict, List
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from oauth2client.service_account import ServiceAccountCredentials

logger = logging.getLogger(__name__)


class DocumentsMerger:
    def __init__(self, gsheets_manager, credentials_file: str, spreadsheet_id: str):
        self.gsheets = gsheets_manager
        self.credentials_file = credentials_file
        self.spreadsheet_id = spreadsheet_id
        self.output_dir = "merged_documents"
        self.docs_service = None
        self._init_services()

        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)

    def _init_services(self):
        """Инициализация Google Docs API"""
        try:
            scope = ['https://www.googleapis.com/auth/documents.readonly']
            creds = ServiceAccountCredentials.from_json_keyfile_name(
                self.credentials_file, scope)
            self.docs_service = build('docs', 'v1', credentials=creds)
            logger.info("Google Docs API инициализирован для чтения")
        except Exception as e:
            logger.error(f"Ошибка инициализации Google Docs API: {e}")

    def merge_qualification_documents(self, target_date: str) -> str:
        """Объединяет все документы квалификаций в один файл"""
        try:
            logger.info("Начинаем объединение документов...")

            # Создаем основной документ
            main_doc = Document()

            # Заголовок
            title = main_doc.add_heading('Объединенные материалы по предметам', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            main_doc.add_paragraph(f"Дата: {target_date}")
            main_doc.add_paragraph(f"Сгенерировано: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
            main_doc.add_paragraph()
            main_doc.add_paragraph("=" * 80)
            main_doc.add_paragraph()

            # Получаем все уникальные ссылки на документы
            document_links = self._get_all_document_links()

            if not document_links:
                main_doc.add_paragraph("❌ Не найдено документов для объединения")
                main_doc.add_paragraph("Проверьте ссылки в листе 'Предметы бот'")
            else:
                # Обрабатываем каждый документ
                for i, (subject_name, doc_url) in enumerate(document_links.items(), 1):
                    success = self._add_document_content(main_doc, subject_name, doc_url, i)
                    if not success:
                        continue

            # Сохраняем объединенный файл
            filename = f"merged_materials_{target_date.replace('.', '_')}.docx"
            filepath = os.path.join(self.output_dir, filename)
            main_doc.save(filepath)

            abs_path = os.path.abspath(filepath)
            logger.info(f"Объединенный документ создан: {abs_path}")

            return f"✅ Документы объединены!\n📁 {abs_path}\n\nОбъединено документов: {len(document_links)}"

        except Exception as e:
            logger.error(f"Ошибка объединения документов: {e}")
            return f"❌ Ошибка объединения: {str(e)}"

    def _get_all_document_links(self) -> Dict[str, str]:
        """Получает все уникальные ссылки на документы из квалификаций"""
        try:
            document_links = {}

            # Получаем данные из листа "Предметы бот"
            worksheet = self.gsheets._get_or_create_worksheet("Предметы бот")
            data = worksheet.get_all_values()

            if len(data) < 2:
                logger.warning("В листе 'Предметы бот' нет данных")
                return document_links

            # Пропускаем заголовок
            for row in data[1:]:
                if len(row) >= 3:  # Есть ссылка в колонке C
                    subject_name = row[1].strip() if len(row) > 1 else "Без названия"
                    doc_url = row[2].strip() if len(row) > 2 else ""

                    if doc_url and doc_url.startswith('http'):
                        document_links[subject_name] = doc_url
                        logger.info(f"Найдена ссылка для предмета '{subject_name}': {doc_url}")

            logger.info(f"Всего найдено ссылок: {len(document_links)}")
            return document_links

        except Exception as e:
            logger.error(f"Ошибка получения ссылок: {e}")
            return {}

    def _add_document_content(self, main_doc: Document, subject_name: str, doc_url: str, index: int) -> bool:
        """Добавляет содержимое одного документа в основной"""
        try:
            # Извлекаем ID документа из URL
            doc_id = self._extract_doc_id(doc_url)
            if not doc_id:
                logger.warning(f"Не удалось извлечь ID из URL: {doc_url}")
                return False

            logger.info(f"Получаем содержимое документа {index}: {subject_name}")

            # Получаем содержимое документа
            content = self._get_document_content(doc_id)
            if not content:
                logger.warning(f"Не удалось получить содержимое для {subject_name}")
                return False

            # Добавляем раздел с содержимым
            main_doc.add_heading(f'{index}. {subject_name}', level=1)

            # Добавляем ссылку на оригинал
            link_para = main_doc.add_paragraph()
            link_para.add_run('📎 Оригинал: ').bold = True
            link_para.add_run(doc_url)

            # Добавляем содержимое
            content_para = main_doc.add_paragraph()
            content_para.add_run(content)

            main_doc.add_paragraph()
            main_doc.add_paragraph("―" * 60)
            main_doc.add_paragraph()

            logger.info(f"Успешно добавлен документ: {subject_name}")
            return True

        except Exception as e:
            logger.error(f"Ошибка добавления документа {subject_name}: {e}")
            return False

    def _extract_doc_id(self, url: str) -> str:
        """Извлекает ID документа из URL"""
        try:
            if '/d/' in url:
                return url.split('/d/')[1].split('/')[0]
            elif 'id=' in url:
                return url.split('id=')[1].split('&')[0]
            return ""
        except Exception as e:
            logger.error(f"Ошибка извлечения ID из {url}: {e}")
            return ""

    def _get_document_content(self, doc_id: str) -> str:
        """Получает текстовое содержимое Google документа"""
        try:
            document = self.docs_service.documents().get(documentId=doc_id).execute()
            content_parts = []

            # Обрабатываем все элементы документа
            for element in document.get('body', {}).get('content', []):
                if 'paragraph' in element:
                    paragraph = element['paragraph']
                    paragraph_text = self._extract_paragraph_text(paragraph)
                    if paragraph_text:
                        content_parts.append(paragraph_text)
                elif 'table' in element:
                    # Для таблиц добавляем маркер
                    content_parts.append("[Таблица]")

            full_content = '\n'.join(content_parts)
            logger.info(f"Получено содержимое документа {doc_id}: {len(full_content)} символов")
            return full_content

        except HttpError as e:
            logger.error(f"Ошибка доступа к документу {doc_id}: {e}")
            return f"[Ошибка доступа к документу: {e}]"
        except Exception as e:
            logger.error(f"Ошибка получения содержимого {doc_id}: {e}")
            return f"[Ошибка: {str(e)}]"

    def _extract_paragraph_text(self, paragraph: Dict) -> str:
        """Извлекает текст из параграфа"""
        try:
            text_parts = []
            for elem in paragraph.get('elements', []):
                if 'textRun' in elem:
                    text_content = elem['textRun'].get('content', '')
                    text_parts.append(text_content)

            text = ''.join(text_parts).strip()
            return text if text else ""
        except Exception as e:
            logger.error(f"Ошибка извлечения текста параграфа: {e}")
            return ""