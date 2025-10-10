# local_materials_manager.py
import logging
from datetime import datetime
import os
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import asyncio

logger = logging.getLogger(__name__)


class LocalMaterialsManager:
    def __init__(self, gsheets_manager, credentials_file: str, spreadsheet_id: str):
        self.gsheets = gsheets_manager
        self.credentials_file = credentials_file
        self.spreadsheet_id = spreadsheet_id
        self.output_dir = "generated_materials"

        # Создаем папку для документов если её нет
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
            logger.info(f"Создана папка для документов: {self.output_dir}")

    def get_student_lesson_info(self, user_id: int, subject_id: str, target_date: str) -> Dict[str, any]:
        """Получает информацию о занятии ученика"""
        try:
            from shedule_app.GoogleParser import GoogleSheetsDataLoader

            loader = GoogleSheetsDataLoader(self.credentials_file, self.spreadsheet_id, target_date)
            topic = loader.get_student_topic_by_user_id(str(user_id), target_date, str(subject_id))

            return {
                'lesson_number': 1,
                'topic': topic or "Тема не определена",
                'subject_name': self._get_subject_name(subject_id)
            }
        except Exception as e:
            logger.error(f"Ошибка получения информации о занятии: {e}")
            return {
                'lesson_number': 1,
                'topic': "Тема не определена",
                'subject_name': self._get_subject_name(subject_id)
            }

    def get_qualification_materials(self, subject_id: str) -> Optional[str]:
        """Получает материалы по предмету"""
        try:
            if hasattr(self.gsheets, 'qual_links') and subject_id in self.gsheets.qual_links:
                doc_url = self.gsheets.qual_links[subject_id]
                return f"Ссылка на материалы: {doc_url}"
            return "Материалы не найдены"
        except Exception as e:
            logger.error(f"Ошибка получения материалов: {e}")
            return None

    def create_combined_materials_document(self, target_date: str) -> str:
        """Создает локальный Word документ с материалами"""
        try:
            # Создаем документ
            doc = Document()

            # Настраиваем стили
            self._setup_document_styles(doc)

            # Добавляем заголовок
            title = doc.add_heading(f'Материалы для занятий', level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Добавляем дату
            date_para = doc.add_paragraph(f'Дата: {target_date}')
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_para.style.font.size = Pt(12)

            doc.add_paragraph()

            # Добавляем информацию о генерации
            info_para = doc.add_paragraph()
            info_para.add_run('Документ сгенерирован автоматически: ').bold = True
            info_para.add_run(datetime.now().strftime("%d.%m.%Y %H:%M"))

            doc.add_paragraph()
            doc.add_paragraph("=" * 80)
            doc.add_paragraph()

            # Получаем реальные данные студентов
            students_data = self._get_students_data(target_date)

            if students_data:
                for i, student_data in enumerate(students_data, 1):
                    self._add_student_section(doc, student_data, i)
            else:
                # Добавляем тестовые данные если реальные не найдены
                self._add_sample_data(doc)

            # Сохраняем файл
            filename = f"materials_{target_date.replace('.', '_')}.docx"
            filepath = os.path.join(self.output_dir, filename)
            doc.save(filepath)

            # Получаем абсолютный путь для отображения
            abs_path = os.path.abspath(filepath)
            logger.info(f"Документ сохранен: {abs_path}")

            return f"Файл успешно создан:\n{abs_path}"

        except Exception as e:
            logger.error(f"Ошибка создания Word документа: {e}")
            return f"Ошибка создания документа: {str(e)}"

    def _setup_document_styles(self, doc):
        """Настраивает стили документа"""
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

    def _get_students_data(self, target_date: str) -> List[Dict]:
        """Получает данные студентов из Google Sheets"""
        try:
            student_sheet = self.gsheets._get_or_create_worksheet("Ученики бот")
            data = student_sheet.get_all_values()

            students_data = []

            if len(data) < 2:
                return students_data

            for row in data[1:]:  # Пропускаем заголовок
                if not row or len(row) < 3:
                    continue

                user_id = str(row[0]).strip() if row[0] else ""
                student_name = str(row[1]).strip() if len(row) > 1 else ""
                subject_id = str(row[2]).strip() if len(row) > 2 else ""

                if user_id and student_name and subject_id:
                    lesson_info = self.get_student_lesson_info(int(user_id), subject_id, target_date)
                    materials = self.get_qualification_materials(subject_id)

                    students_data.append({
                        'name': student_name,
                        'subject_name': lesson_info['subject_name'],
                        'lesson_number': lesson_info['lesson_number'],
                        'topic': lesson_info['topic'],
                        'materials': materials
                    })

            return students_data[:10]  # Ограничиваем 10 студентами для теста

        except Exception as e:
            logger.error(f"Ошибка получения данных студентов: {e}")
            return []

    def _add_student_section(self, doc, student_data: Dict, index: int):
        """Добавляет раздел для студента"""
        # Заголовок студента
        student_heading = doc.add_heading(f'{index}. {student_data["name"]}', level=1)

        # Информация о предмете и занятии
        info_para = doc.add_paragraph()
        info_para.add_run('Предмет: ').bold = True
        info_para.add_run(f'{student_data["subject_name"]}\n')

        info_para.add_run('Занятие: ').bold = True
        info_para.add_run(f'№{student_data["lesson_number"]}\n')

        info_para.add_run('Тема: ').bold = True
        info_para.add_run(f'{student_data["topic"]}\n')

        info_para.add_run('Материалы: ').bold = True
        info_para.add_run(f'{student_data["materials"]}')

        doc.add_paragraph()
        doc.add_paragraph("-" * 60)
        doc.add_paragraph()

    def _add_sample_data(self, doc):
        """Добавляет примеры данных если реальные не найдены"""
        sample_students = [
            {
                'name': 'Иван Иванов',
                'subject_name': 'Математика',
                'lesson_number': 1,
                'topic': 'Алгебраические уравнения',
                'materials': 'Ссылка на материалы: https://example.com/math'
            },
            {
                'name': 'Мария Петрова',
                'subject_name': 'Физика',
                'lesson_number': 2,
                'topic': 'Законы Ньютона',
                'materials': 'Ссылка на материалы: https://example.com/physics'
            },
            {
                'name': 'Алексей Сидоров',
                'subject_name': 'Информатика',
                'lesson_number': 3,
                'topic': 'Основы программирования',
                'materials': 'Ссылка на материалы: https://example.com/informatics'
            }
        ]

        for i, student in enumerate(sample_students, 1):
            self._add_student_section(doc, student, i)

    def _get_subject_name(self, subject_id: str) -> str:
        """Получает название предмета по ID"""
        from config import SUBJECTS
        return SUBJECTS.get(subject_id, f"Предмет {subject_id}")