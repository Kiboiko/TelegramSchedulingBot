import os
import io
import requests
import tempfile
from typing import List
from docx import Document
from docx.shared import Inches
import gspread
from oauth2client.service_account import ServiceAccountCredentials
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
import logging
from bs4 import BeautifulSoup
import re

logger = logging.getLogger(__name__)

class GoogleDocsMerger:
    def __init__(self, credentials_file: str):
        """
        Инициализирует класс для работы с Google Docs
        
        Args:
            credentials_file: путь к файлу учетных данных Google Service Account
        """
        self.credentials_file = credentials_file
        self.client = None
        self.docs_service = None
        self.drive_service = None
        self._connect()
    
    def _connect(self):
        """Устанавливает соединение с Google Docs API"""
        try:
            scope = [
                'https://www.googleapis.com/auth/documents',
                'https://www.googleapis.com/auth/drive',
                'https://spreadsheets.google.com/feeds',
                'https://www.googleapis.com/auth/drive.readonly'
            ]
            creds = ServiceAccountCredentials.from_json_keyfile_name(
                self.credentials_file, scope)
            
            # Авторизация для gspread
            self.client = gspread.authorize(creds)
            
            # Создаем отдельный сервис для Google Docs API
            self.docs_service = build('docs', 'v1', credentials=creds)
            self.drive_service = build('drive', 'v3', credentials=creds)
            
            logger.info("Успешное подключение к Google APIs")
            return True
        except Exception as e:
            logger.error(f"Ошибка подключения к Google APIs: {e}")
            return False

    def check_document_access(self, doc_url: str) -> bool:
        """Проверяет доступ к документу"""
        try:
            doc_id = self.extract_doc_id_from_url(doc_url)
            if not doc_id:
                return False
            
            # Пробуем получить метаданные документа
            document = self.docs_service.documents().get(documentId=doc_id).execute()
            logger.info(f"Доступ к документу есть: {document.get('title', 'Unknown')}")
            return True
        except HttpError as e:
            logger.error(f"Ошибка доступа: {e}")
            return False
    
    def extract_doc_id_from_url(self, url: str) -> str:
        """
        Извлекает ID документа из URL Google Docs
        """
        try:
            if '/document/d/' in url:
                doc_id = url.split('/document/d/')[1].split('/')[0]
            elif 'id=' in url:
                doc_id = url.split('id=')[1].split('&')[0]
            else:
                doc_id = url.split('/')[-1].split('?')[0]
            
            return doc_id
        except Exception as e:
            logger.error(f"Ошибка извлечения ID из URL {url}: {e}")
            return None
    
    def get_document_content_with_images(self, doc_url: str) -> dict:
        """
        Получает содержимое Google Doc с изображениями
        """
        try:
            doc_id = self.extract_doc_id_from_url(doc_url)
            if not doc_id:
                return {'text': '', 'images': []}
            
            # Сначала пробуем через HTML экспорт
            logger.info("Пробуем метод HTML экспорта...")
            result = self.get_document_content_with_images_improved(doc_url)
            
            if result['images']:
                logger.info(f"Найдено {len(result['images'])} изображений через HTML экспорт")
                return result
            
            # Если изображений нет, пробуем через Drive API
            logger.info("Пробуем метод Drive API...")
            images_via_drive = self.get_images_via_drive_api(doc_id)
            if images_via_drive:
                text_content = self._get_document_content_simple(doc_url)
                return {'text': text_content, 'images': images_via_drive}
            
            # Если все методы не сработали, возвращаем только текст
            logger.info("Изображения не найдены, возвращаем только текст")
            text_content = self._get_document_content_simple(doc_url)
            return {'text': text_content, 'images': []}
            
        except Exception as e:
            logger.error(f"Ошибка получения содержимого документа: {e}")
            return {'text': '', 'images': []}

    def get_document_content_with_images_improved(self, doc_url: str) -> dict:
        """
        Улучшенный метод получения содержимого с изображениями через экспорт HTML
        """
        try:
            doc_id = self.extract_doc_id_from_url(doc_url)
            if not doc_id:
                return {'text': '', 'images': []}

            # Экспортируем документ в HTML
            export_url = f"https://docs.google.com/document/d/{doc_id}/export?format=html"
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            response = requests.get(export_url, headers=headers, timeout=30)
            
            if response.status_code != 200:
                logger.error(f"Ошибка экспорта HTML: {response.status_code}")
                return {'text': '', 'images': []}

            html_content = response.text
            text_content, images_data = self._parse_html_content(html_content, doc_id)
            
            return {'text': text_content, 'images': images_data}
            
        except Exception as e:
            logger.error(f"Ошибка получения содержимого: {e}")
            return {'text': '', 'images': []}

    def _parse_html_content(self, html_content: str, doc_id: str):
        """Парсит HTML контент и извлекает текст и изображения"""
        try:
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Удаляем ненужные теги
            for element in soup(['script', 'style', 'meta', 'link']):
                element.decompose()
            
            # Извлекаем текст
            text_content = soup.get_text(separator='\n', strip=True)
            text_content = re.sub(r'\n\s*\n', '\n\n', text_content)  # Убираем лишние переносы
            
            # Извлекаем изображения
            images_data = []
            img_tags = soup.find_all('img')
            
            for i, img_tag in enumerate(img_tags):
                try:
                    img_src = img_tag.get('src', '')
                    if not img_src:
                        continue
                        
                    # Скачиваем изображение
                    image_data = self._download_image_from_url(img_src)
                    if image_data:
                        images_data.append({
                            'data': image_data,
                            'index': i,
                            'alt': img_tag.get('alt', f'Image {i+1}'),
                            'width': img_tag.get('width', '400'),
                            'height': img_tag.get('height', '300')
                        })
                        logger.info(f"Изображение {i+1} успешно загружено")
                        
                except Exception as e:
                    logger.error(f"Ошибка обработки изображения {i}: {e}")
                    continue
            
            return text_content, images_data
            
        except Exception as e:
            logger.error(f"Ошибка парсинга HTML: {e}")
            return "", []

    def _download_image_from_url(self, image_url: str) -> bytes:
        """Скачивает изображение по URL"""
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
                'Accept': 'image/webp,image/apng,image/*,*/*',
                'Accept-Language': 'en-US,en;q=0.9',
            }
            
            response = requests.get(image_url, headers=headers, timeout=30, stream=True)
            if response.status_code == 200:
                return response.content
            else:
                logger.warning(f"Не удалось загрузить изображение: {response.status_code}")
        except Exception as e:
            logger.error(f"Ошибка загрузки изображения: {e}")
        return None

    def get_images_via_drive_api(self, doc_id: str):
        """Получает изображения через Drive API"""
        try:
            # Получаем список файлов в папке документа
            results = self.drive_service.files().list(
                q=f"'{doc_id}' in parents and mimeType contains 'image/'",
                fields="files(id, name, mimeType, size)"
            ).execute()
            
            images = results.get('files', [])
            images_data = []
            
            logger.info(f"Найдено {len(images)} изображений через Drive API")
            
            for img in images:
                try:
                    # Скачиваем изображение
                    request = self.drive_service.files().get_media(fileId=img['id'])
                    image_content = request.execute()
                    
                    images_data.append({
                        'data': image_content,
                        'name': img['name'],
                        'mime_type': img['mimeType'],
                        'size': img.get('size', 0)
                    })
                    
                    logger.info(f"Загружено изображение: {img['name']}")
                    
                except Exception as e:
                    logger.error(f"Ошибка загрузки изображения {img['name']}: {e}")
                    continue
                    
            return images_data
            
        except Exception as e:
            logger.error(f"Ошибка Drive API: {e}")
            return []

    def _get_document_content_simple(self, doc_url: str) -> str:
        """
        Упрощенный метод получения содержимого через экспорт
        """
        try:
            doc_id = self.extract_doc_id_from_url(doc_url)
            if not doc_id:
                return ""
            
            # Пробуем разные форматы экспорта
            formats = ['txt', 'html']
            
            for format_type in formats:
                try:
                    export_url = f"https://docs.google.com/document/d/{doc_id}/export?format={format_type}"
                    headers = {
                        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
                    }
                    response = requests.get(export_url, headers=headers, timeout=30)
                    
                    if response.status_code == 200:
                        if format_type == 'html':
                            # Для HTML извлекаем чистый текст
                            return self._extract_text_from_html(response.text)
                        else:
                            return response.text
                except Exception as e:
                    logger.warning(f"Не удалось экспортировать в формате {format_type}: {e}")
                    continue
            
            return ""
            
        except Exception as e:
            logger.error(f"Ошибка упрощенного получения содержимого: {e}")
            return ""

    def _extract_text_from_html(self, html_content):
        """Извлекает текст из HTML"""
        try:
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Удаляем ненужные теги
            for element in soup(['script', 'style', 'meta', 'link']):
                element.decompose()
            
            # Получаем чистый текст
            text = soup.get_text(separator='\n', strip=True)
            text = re.sub(r'\n\s*\n', '\n\n', text)  # Убираем лишние переносы
            
            return text
        except Exception as e:
            logger.error(f"Ошибка извлечения текста из HTML: {e}")
            # Возвращаем текст без обработки если что-то пошло не так
            clean = re.compile('<.*?>')
            return re.sub(clean, '', html_content)
    
    def merge_documents_with_images(self, doc_urls: List[str], output_path: str) -> bool:
        """
        Объединяет несколько Google Docs в один Word-документ с попыткой сохранения изображений
        """
        try:
            doc = Document()
            title = doc.add_heading('Объединенный документ', 0)
            doc.add_paragraph(f"Объединено документов: {len(doc_urls)}")
            doc.add_paragraph()
            
            total_images = 0
            
            for i, doc_url in enumerate(doc_urls, 1):
                logger.info(f"Обработка документа {i}/{len(doc_urls)}: {doc_url}")
                
                if i > 1:
                    doc.add_paragraph("\n" + "="*50 + "\n")
                
                doc.add_heading(f'Документ {i}', level=1)
                
                # Получаем контент с информацией об изображениях
                content_data = self.get_document_content_with_images(doc_url)
                text_content = content_data['text']
                images = content_data['images']
                
                total_images += len(images)
                logger.info(f"Документ {i}: {len(text_content)} символов, {len(images)} изображений")
                
                if text_content:
                    # Обрабатываем текст и добавляем изображения
                    paragraphs = text_content.split('\n')
                    
                    for paragraph in paragraphs:
                        if paragraph.strip():
                            doc.add_paragraph(paragraph.strip())
                    
                    # Добавляем изображения в конец документа
                    if images:
                        doc.add_paragraph("\n" + "-" * 30 + " ИЗОБРАЖЕНИЯ " + "-" * 30 + "\n")
                        
                        for img_idx, image_info in enumerate(images, 1):
                            try:
                                image_data = image_info['data']
                                if image_data:
                                    # Создаем временный файл для изображения
                                    image_stream = io.BytesIO(image_data)
                                    
                                    # Добавляем подпись к изображению
                                    doc.add_paragraph(f"Изображение {img_idx}: {image_info.get('alt', '')}")
                                    
                                    # Добавляем изображение
                                    doc.add_picture(image_stream, width=Inches(5.0))
                                    
                                    # Пустая строка после изображения
                                    doc.add_paragraph()
                                    
                                    logger.info(f"Добавлено изображение {img_idx} в документ {i}")
                                    
                            except Exception as e:
                                logger.error(f"Ошибка добавления изображения {img_idx}: {e}")
                                doc.add_paragraph(f"[Ошибка загрузки изображения {img_idx}]")
                else:
                    doc.add_paragraph(f"Не удалось загрузить содержимое документа {i}")
                    logger.warning(f"Не удалось загрузить документ {i}")
            
            doc.save(output_path)
            logger.info(f"Успешно создан объединенный документ: {output_path}")
            logger.info(f"Всего обработано изображений: {total_images}")
            return True
            
        except Exception as e:
            logger.error(f"Ошибка объединения документов: {e}")
            return False

    # Старый метод для обратной совместимости
    def merge_documents(self, doc_urls: List[str], output_path: str) -> bool:
        return self.merge_documents_with_images(doc_urls, output_path)


# Альтернативный подход: экспорт в PDF и конвертация
def export_as_pdf_and_convert(doc_url: str, output_path: str) -> bool:
    """
    Альтернативный метод: экспортирует документ как PDF и конвертирует в Word
    Это может сохранить изображения, но требует дополнительных библиотек
    """
    try:
        import requests
        from pdf2docx import Converter
        
        doc_id = doc_url.split('/document/d/')[1].split('/')[0]
        pdf_url = f"https://docs.google.com/document/d/{doc_id}/export?format=pdf"
        
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        
        response = requests.get(pdf_url, headers=headers, timeout=30)
        if response.status_code == 200:
            # Сохраняем временный PDF
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_pdf:
                temp_pdf.write(response.content)
                temp_pdf_path = temp_pdf.name
            
            try:
                # Конвертируем PDF в Word
                cv = Converter(temp_pdf_path)
                cv.convert(output_path)
                cv.close()
                
                logger.info(f"Успешно конвертировано через PDF: {output_path}")
                return True
            finally:
                # Удаляем временный файл
                os.unlink(temp_pdf_path)
        
        logger.error(f"Ошибка экспорта PDF: {response.status_code}")
        return False
        
    except ImportError:
        logger.error("Библиотека pdf2docx не установлена. Установите: pip install pdf2docx")
        return False
    except Exception as e:
        logger.error(f"Ошибка экспорта через PDF: {e}")
        return False


# Улучшенная функция проверки доступа
def check_document_permissions(merger: GoogleDocsMerger, doc_urls: List[str]):
    """Проверяет доступ ко всем документам"""
    for i, doc_url in enumerate(doc_urls, 1):
        logger.info(f"Проверка доступа к документу {i}: {doc_url}")
        has_access = merger.check_document_access(doc_url)
        
        if has_access:
            logger.info(f"✅ Доступ к документу {i} ЕСТЬ")
        else:
            logger.warning(f"❌ Доступа к документу {i} НЕТ")
            
            # Даем подсказку по настройке доступа
            doc_id = merger.extract_doc_id_from_url(doc_url)
            if doc_id:
                logger.info(f"💡 Для предоставления доступа:")
                logger.info(f"1. Откройте настройки доступа документа")
                logger.info(f"2. Добавьте email сервисного аккаунта как редактора")
                logger.info(f"3. Или откройте доступ 'Все, у кого есть ссылка'")


# Пример использования
# def main():
#     merger = GoogleDocsMerger('credentials.json')
    
#     doc_urls = [
#         'https://docs.google.com/document/d/1009vY3SSX_BUyU8wGD9YQrFETRB5o34FNBmgOA-VAwk/edit?usp=sharing',
#     ]
    
#     # Проверяем доступ
#     check_document_permissions(merger, doc_urls)
    
#     # Основной метод с изображениями
#     success = merger.merge_documents_with_images(doc_urls, 'объединенный_документ.docx')
    
#     if not success:
#         # Альтернативный метод через PDF
#         print("Пробуем альтернативный метод через PDF...")
#         success = export_as_pdf_and_convert(doc_urls[0], 'объединенный_документ_альтернативный.docx')
    
#     if success:
#         print("✅ Документы успешно объединены!")
#     else:
#         print("❌ Произошла ошибка при объединении документов")

# if __name__ == "__main__":
#     main()