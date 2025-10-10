# alternative_materials_manager.py
import logging
from datetime import datetime
import os
from docx import Document
from docx.shared import Inches

logger = logging.getLogger(__name__)


class LocalMaterialsManager:
    def __init__(self, gsheets_manager, credentials_file: str, spreadsheet_id: str):
        self.gsheets = gsheets_manager
        self.credentials_file = credentials_file
        self.spreadsheet_id = spreadsheet_id
        self.output_dir = "materials_documents"

        # Создаем папку для документов если её нет
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)

    def create_combined_materials_document(self, target_date: str) -> str:
        """Создает локальный Word документ с материалами"""
        try:
            # Создаем документ
            doc = Document()

            # Добавляем заголовок
            title = doc.add_heading(f'Материалы для занятий на {target_date}', 0)

            # Добавляем дату генерации
            doc.add_paragraph(f'Документ сгенерирован: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
            doc.add_paragraph()

            # Добавляем тестовое содержимое
            doc.add_heading('Пример ученика 1:', level=1)
            p1 = doc.add_paragraph()
            p1.add_run('🎓 Ученик: ').bold = True
            p1.add_run('Иван Иванов\n')
            p1.add_run('📖 Предмет: ').bold = True
            p1.add_run('Математика\n')
            p1.add_run('🔢 Занятие: ').bold = True
            p1.add_run('№1\n')
            p1.add_run('📌 Тема: ').bold = True
            p1.add_run('Алгебраические уравнения\n')
            p1.add_run('📝 Материалы: ').bold = True
            p1.add_run('https://example.com/math')

            doc.add_paragraph()
            doc.add_paragraph('=' * 50)
            doc.add_paragraph()

            doc.add_heading('Пример ученика 2:', level=1)
            p2 = doc.add_paragraph()
            p2.add_run('🎓 Ученик: ').bold = True
            p2.add_run('Мария Петрова\n')
            p2.add_run('📖 Предмет: ').bold = True
            p2.add_run('Физика\n')
            p2.add_run('🔢 Занятие: ').bold = True
            p2.add_run('№2\n')
            p2.add_run('📌 Тема: ').bold = True
            p2.add_run('Законы Ньютона\n')
            p2.add_run('📝 Материалы: ').bold = True
            p2.add_run('https://example.com/physics')

            # Сохраняем файл
            filename = f"materials_{target_date.replace('.', '_')}.docx"
            filepath = os.path.join(self.output_dir, filename)
            doc.save(filepath)

            logger.info(f"Документ сохранен: {filepath}")
            return f"Файл сохранен: {filepath}"

        except Exception as e:
            logger.error(f"Ошибка создания Word документа: {e}")
            return f"Ошибка: {str(e)}"